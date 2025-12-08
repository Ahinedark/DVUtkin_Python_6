"""Скрипт, генерирующий отчет о структуре файлов и папок на жестком диске.

1. Путь к анализируемой папке задаётся в параметре командной строки --path (обязательный);
2. Путь к отчету задаётся через параметр командной строки --report (необязательный, default = report.json)
3. Тип отчета (docx, xlsx, pdf, csv, json) определяется из расширения файла отчета. Если указано другое расширение, выдаётся сообщение об ошибке;
4. Формат вывода в каждый вид отчета - дерево. Выводится имя, размер (для файлов, для папок - слово FOLDER, для ZIP-архивов - слово ZIP), дата изменения
5. Анализируется текущая папка и вложенные папки. Если встречаются ZIP-архивы, они рассматриваются как вложенные папки. Символические ссылки не анализируются
"""

import argparse
import csv
import json
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Indenter, Paragraph, SimpleDocTemplate, Spacer


class Item:
    """Класс для элемента отчёта (объект: файл, папка или архив).

    Атрибуты:
        path (str): Путь к объекту в виде строки.
        size (int): Размер объекта (для файла, для папки - слово FOLDER, для ZIP-архивов - слово ZIP).
        time (datetime): Время изменения.
        level (int): Уровень вложенности объекта (для красивого вывода).
    """

    def __init__(self, root_path:str, path:str, size:int, time:datetime):
        """Конструктор класса Item.

        Args:
            root_path (str): Путь к анализируемой скриптом папке (из параметра --path командной строки).
            path (str): Путь к файлу/папке в виде строки.
            size (int): Размер файла (для файла, для папки - слово FOLDER, для ZIP-архивов - слово ZIP).
            time (datetime): Время изменения.
        """
        self.path = str(path)
        self.level = self.path.count('\\') - str(root_path).count('\\') - 1
        # Имя достаём из пути после последнего знака '\'
        self.name = self.path.split('\\')[-1]
        self.size = str(size)
        self.time = str(time)

    def __str__(self):
        """Строковое представление экземпляра Item.

        Returns:
            str: Строка формата "уровень {имя файла} {размер файла} {дата изменения}".
        """
        result = ' ' * (self.level*5) + '' + self.name
        result += ' ' * 5 + '' + self.size
        result += ' ' * 5 + '' + self.time
        return result


def analyze_zip(root:Path, zipf: zipfile.ZipFile, zip_root:Path) -> list[Item]:
    """Функция для анализа zip-архива.

    Args:
        root (Path): Путь к анализируемой скриптом папке (из параметра --path командной строки).
        zipf (zipfile.ZipFile): Открытый текущий zip-архив.
        zip_root (Path): Путь к текущему zip-архиву.

    Returns:
        list[Item]: Список элементов отчёта (объект: файл, папка или архив). 
    """
    items = []
    
    # Проходимся по всем объектам в zip-архиве и добавляем их в отчёт
    for f_name in zipf.infolist():
        path = zip_root.joinpath(f_name.filename)
        size = f_name.file_size
        time = f_name.date_time

        # Если имя файла (из infolist()) заказчивается на '/' - это папка
        if f_name.filename.endswith('/'):
            size = 'FOLDER'

        # Если имя файла заканчивается на '.zip' - это zip-архив
        if path.suffix.lower() == '.zip':
            size = 'ZIP'
            # Открываем и читаем вложенный zip-архив
            with zipf.open(f_name) as inner_file:
                data = inner_file.read()
                with zipfile.ZipFile(BytesIO(data)) as nested_zip:
                    # Запускаем analyze_zip рекурсивно для анализа вложенного zip-архива
                    items.extend(analyze_zip(root, nested_zip, path))
        
        # Добавляем объект в отчёт
        items.append(Item(str(root), str(path), size, time))
    return items


def analyze_folder(root:Path) -> list[Item]:
    """Функция для анализа папки.

    Args:
        root (Path): Путь к анализируемой скриптом папке (из параметра --path командной строки).

    Returns:
        list[Item]: Список элементов отчёта (объект: файл, папка или архив). 
    """
    items = []

    # Рекурсивный просмотр всех объектов в папке
    for path in root.rglob('*'):
        # Игнорируем символические ссылки
        if path.is_symlink():
            continue

        size = path.stat().st_size
        time = datetime.fromtimestamp(path.stat().st_mtime).isoformat()
        
        # Если объект - папка, записываем ему size = 'FOLDER'
        if path.is_dir():
            size = 'FOLDER'
        
        # Если объект - zip-архив, записываем ему size = 'ZIP'
        if path.suffix.lower() == '.zip':
            size = 'ZIP'
            # Открываем и читаем  zip-архив
            with zipfile.ZipFile(path, 'r') as zipf:
                items.extend(analyze_zip(root, zipf, path))

        # Добавляем объект в отчёт
        items.append(Item(str(root), str(path), size, time))        
    
    # Сортируем отчёт по пути (для исключения перемешивания c отчётоv из analyze_zip)
    items.sort(key=lambda x: x.path.lower())
    return items


def write_docx(folder_path:Path, report_path:Path, data:list[Item]):
    """Запись отчёта в документ MS Word (.docx).

    Args:
        folder_path (Path): Путь к анализируемой скриптом папке (из параметра --path командной строки).
        report_path (Path): Путь к файлу отчёта (из параметра --report командной строки).
        data (list[Item]): Список элементов отчёта (объект: файл, папка или архив).
    """
    doc = Document()
    doc.add_heading(f'Отчет о структуре файлов в папке {str(folder_path)}', 1)

    # Выводим элементы отчёта в стрковом представлении
    for d in data:
        doc.add_paragraph(str(d))

    doc.save(report_path)


def write_xlsx(folder_path:Path, report_path:Path, data:list[Item]):
    """Запись отчёта в документ MS Excel (.xlsx).

    Args:
        folder_path (Path): Путь к анализируемой скриптом папке (из параметра --path командной строки).
        report_path (Path): Путь к файлу отчёта (из параметра --report командной строки).
        data (list[Item]): Список элементов отчёта (объект: файл, папка или архив).
    """
    wb = Workbook()
    ws = wb.active

    # Пишем заголовок в объёдинённых ячейках
    ws['A1'] = f'Отчет о структуре файлов в папке {str(folder_path)}'
    ws.merge_cells('A1:E1')
    ws.append([])

    # Выводим элементы отчёта с учётом уровня вложенности
    for d in data:
        ws.append([''] * d.level + [d.name, d.size, d.time])
    
    wb.save(report_path)


def write_pdf(folder_path:Path, report_path:Path, data:list[Item]):
    """Запись отчёта в документ PDF (.pdf).

    Args:
        folder_path (Path): Путь к анализируемой скриптом папке (из параметра --path командной строки).
        report_path (Path): Путь к файлу отчёта (из параметра --report командной строки).
        data (list[Item]): Список элементов отчёта (объект: файл, папка или архив).
    """
    pdf = SimpleDocTemplate(str(report_path), pagesize=A4)
    styles = getSampleStyleSheet()

    # Регистрируем шрифт DejaVuSerif для корректного вывода кириллицйы
    styles['Normal'].fontName='DejaVuSerif'
    styles['Title'].fontName='DejaVuSerif'
    # Файл DejaVuSerif.ttf должен быть обязательно
    pdfmetrics.registerFont(TTFont('DejaVuSerif','DejaVuSerif.ttf', 'UTF-8'))

    elements = []

    title = Paragraph(f'Анализ папки {folder_path}', styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 20))

    # Выводим элементы отчёта с учётом уровня вложенности
    for d in data:
        indent = 20 * d.level
        elements.append(Indenter(left=indent))
        line = Paragraph(str(d), styles['Normal'])
        elements.append(line)
        elements.append(Spacer(1, 4))
        elements.append(Indenter(left=-indent))

    pdf.build(elements)


def write_csv(folder_path:Path, report_path:Path, data:list[Item]):
    """Запись отчёта в файл CSV (.csv).

    Args:
        folder_path (Path): Путь к анализируемой скриптом папке (из параметра --path командной строки).
        report_path (Path): Путь к файлу отчёта (из параметра --report командной строки).
        data (list[Item]): Список элементов отчёта (объект: файл, папка или архив).
    """
    with Path.open(report_path, 'w', encoding='utf-8-sig', newline='') as f:
        writer = csv.writer(f, delimiter=';')
        
        writer.writerow([f'Отчет о структуре файлов в папке {str(folder_path)}'])
        writer.writerow([])

        # Выводим элементы отчёта с учётом уровня вложенности
        for d in data:
            row = [''] * d.level + [d.name, d.size, d.time]
            writer.writerow(row)


def write_json(folder_path:Path, report_path:Path, data:list[Item]):
    """Запись отчёта в файл JSON (.json).

    Args:
        folder_path (Path): Путь к анализируемой скриптом папке (из параметра --path командной строки).
        report_path (Path): Путь к файлу отчёта (из параметра --report командной строки).
        data (list[Item]): Список элементов отчёта (объект: файл, папка или архив).
    """
    json_list = []

    for i in data:
        print(str(i))

    for d in data:
        json_list.append({
            'level': d.level,
            'name': d.name,
            'size': d.size,
            'time': d.time,
            'path': d.path
        })

    output = {
        'folder': str(folder_path),
        'items': json_list
    }

    with Path.open(report_path, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=4)


def main():
    """Главная функция.
    
    Вызывается через командную строку, принимает 2 параметра:
    Путь к анализируемой папке --path (обязательный);
    Путь к отчету --report (необязательный, default = report.json)
    """
    # Парсим параметры из командной строки
    parser = argparse.ArgumentParser(description='Analyze folder')
    parser.add_argument('--path', required=True, type=str, help='Путь к анализируемой папке. Пример: "C:\\Users\\Public"')
    parser.add_argument('--report', required=False, type=str, default='report.json', help='Путь к файлу отчёта. Пример: "C:\\Users\\Public\\report.json"')
    args = parser.parse_args()
    folder_path = Path(args.path)
    report_path = Path(args.report)
    
    # Определяем расширение файла отчёта
    ex = report_path.suffix.lower()

    if not folder_path.is_dir() or not folder_path.exists():
        print('Ошибка: указанный путь не является папкой или не существует.')
        print('Укажите корректный путь к анализируемой папке.')
        return

    # Анализируем папку
    data = analyze_folder(folder_path)

    # Записываем отчёт с указанным расширением
    match ex:
        case '.docx':
            write_docx(folder_path, report_path, data)
        case '.xlsx':
            write_xlsx(folder_path, report_path, data)
        case '.pdf':
            write_pdf(folder_path, report_path, data)
        case '.csv':
            write_csv(folder_path, report_path, data)
        case '.json':
            write_json(folder_path, report_path, data)
        case _:
            print(f'Ошибка: указан некорректный формат отчёта: "{ex}".')
            print('Укажите корректный формат отчёта (".csv", ".json", ".docx", ".xlsx", ".pdf") и попробуйте ещё раз.')
            return

    print('Отчёт сформирован:', str(report_path))

if __name__ == '__main__':
    main()